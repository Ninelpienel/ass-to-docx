import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from tkinter import Tk
from tkinter.filedialog import askopenfilenames

def clean_ass_tags(text):
    text = re.sub(r'\{.*?\}', '', text)
    text = re.sub(r'\\[Nn]', '\n', text)
    text = text.strip()
    return text

def convert_ass_to_docx(file_path):
    doc = Document()
    line_number = 1

    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            if line.startswith('Dialogue:'):
                dialogue = line.split(',', 9)[-1].strip()
                cleaned_text = clean_ass_tags(dialogue)
                cleaned_text = re.sub(r'\n$', '', cleaned_text)
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                tab_stops = paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(0.5), WD_TAB_ALIGNMENT.LEFT)

                run = paragraph.add_run(f"{line_number}:\t")
                parts = cleaned_text.split('\n')
                for i, part in enumerate(parts):
                    if i > 0:
                        run.add_break()
                        run = paragraph.add_run("\t")
                    run.add_text(part)
                line_number += 1

    output_path = os.path.splitext(file_path)[0] + '.docx'
    doc.save(output_path)
    print(f"Converted {file_path} to {output_path}")

def main():
    Tk().withdraw()
    file_paths = askopenfilenames(filetypes=[("ASS files", "*.ass")])
    
    for file_path in file_paths:
        convert_ass_to_docx(file_path)

if __name__ == "__main__":
    main()